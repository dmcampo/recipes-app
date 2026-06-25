# -*- coding: utf-8 -*-
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def generate_document():
    doc = docx.Document()
    
    # Page setup - Standard Letter, 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Configure headers and footers
        footer = section.footer
        f_p = footer.paragraphs[0]
        f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f_run = f_p.add_run("Página ")
        f_run.font.name = 'Calibri'
        f_run.font.size = Pt(9)
        f_run.font.color.rgb = RGBColor(120, 120, 120)
        add_page_number(f_run)

    # Color Palette (Premium Gastronomy theme: Dark Gray, Gold/Amber Accent, Warm Cream)
    COLOR_PRIMARY = RGBColor(26, 27, 30)     # #1A1B1E - Charcoal
    COLOR_SECONDARY = RGBColor(217, 119, 6)  # #D97706 - Amber/Terracotta Accent
    COLOR_MUTED = RGBColor(107, 114, 128)    # #6B7280 - Gray
    COLOR_TEXT = RGBColor(40, 40, 40)        # #282828 - Dark Gray body text

    # Set up basic style definitions
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = COLOR_TEXT

    # ----------------------------------------------------
    # COVER PAGE
    # ----------------------------------------------------
    # Spacer
    for _ in range(3):
        doc.add_paragraph()
        
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("DOCUMENTACIÓN TÉCNICA Y DE ARQUITECTURA\n")
    run_title.font.name = 'Georgia'
    run_title.font.size = Pt(28)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY

    p_subtitle = doc.add_paragraph()
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_subtitle.add_run("Aplicación Web de Recetas Gastronómicas Premium (React + Vite)")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(16)
    run_sub.font.italic = True
    run_sub.font.color.rgb = COLOR_SECONDARY
    
    for _ in range(8):
        doc.add_paragraph()
        
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_runs = [
        ("Desarrollador: ", True), ("Ingeniero de Software Senior & Frontend Expert\n", False),
        ("Tecnologías: ", True), ("React v18, Vite, React Router DOM, Axios, LocalStorage, CSS Puro\n", False),
        ("Fecha: ", True), ("Junio de 2026\n", False),
        ("Estado: ", True), ("Versión Final (Estilo Gastronómico Premium - Producción Aprobada)", False)
    ]
    for text, is_bold in meta_runs:
        run = p_meta.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
        run.font.bold = is_bold
        run.font.color.rgb = COLOR_PRIMARY if is_bold else COLOR_MUTED

    doc.add_page_break()

    # ----------------------------------------------------
    # INTRODUCCIÓN Y ARQUITECTURA GENERAL
    # ----------------------------------------------------
    h1 = doc.add_heading(level=1)
    run_h1 = h1.add_run("1. Introducción y Propósito del Proyecto")
    run_h1.font.name = 'Georgia'
    run_h1.font.size = Pt(18)
    run_h1.font.bold = True
    run_h1.font.color.rgb = COLOR_PRIMARY
    
    p = doc.add_paragraph(
        "Este proyecto consiste en una aplicación web interactiva de recetas de cocina gourmet "
        "diseñada bajo una estética editorial premium (inspirada en sitios de alta cocina como elgourmet.com). "
        "Permite a los usuarios buscar recetas de forma instantánea, filtrar por categorías culinarias, "
        "ver detalles e instrucciones paso a paso, visualizar videotutoriales integrados de YouTube y guardar sus "
        "recetas favoritas con persistencia a nivel de navegador. El diseño sigue principios visuales sofisticados: "
        "sombras suaves, bordes refinados, tipografías elegantes (Playfair Display y Montserrat) y soporte completo "
        "para temas claro y oscuro con una adaptación visual exquisita."
    )
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(12)

    h1 = doc.add_heading(level=1)
    run_h1 = h1.add_run("2. Tecnologías y Librerías Utilizadas")
    run_h1.font.name = 'Georgia'
    run_h1.font.size = Pt(18)
    run_h1.font.bold = True
    run_h1.font.color.rgb = COLOR_PRIMARY

    techs = [
        ("React JS (v18): ", "Biblioteca principal para la construcción de interfaces de usuario interactivas basadas en componentes reactivos y estados controlados."),
        ("Vite JS (v6/8): ", "Entorno de desarrollo rápido y empaquetador moderno para el frontend que optimiza los tiempos de carga en desarrollo y genera compilaciones ultra eficientes en producción."),
        ("React Router DOM (v6): ", "Manejador del enrutamiento SPA (Single Page Application) que gestiona la navegación interna entre Home, Vista de Detalle y Favoritos sin refrescos del navegador."),
        ("Axios: ", "Cliente HTTP para realizar solicitudes asíncronas y optimizadas a la API pública de TheMealDB, proporcionando interceptores implícitos e integración limpia."),
        ("TheMealDB API: ", "API gratuita utilizada para obtener categorías culinarias reales e información completa de recetas (títulos, áreas, instrucciones, ingredientes y videos)."),
        ("CSS Puro (Vanilla CSS): ", "Sistema de estilos implementado mediante variables CSS globales (Custom Properties), garantizando total personalización sin dependencias externas y transiciones fluidas."),
        ("LocalStorage HTML5: ", "Mecanismo de persistencia utilizado en el FavoritesContext para almacenar los identificadores y detalles de las recetas favoritas de forma permanente en el navegador del usuario.")
    ]

    for title, desc in techs:
        p_t = doc.add_paragraph(style='List Bullet')
        p_t.paragraph_format.space_after = Pt(6)
        r_title = p_t.add_run(title)
        r_title.font.bold = True
        r_title.font.color.rgb = COLOR_SECONDARY
        r_desc = p_t.add_run(desc)
        r_desc.font.color.rgb = COLOR_TEXT

    doc.add_paragraph()

    # ----------------------------------------------------
    # ESTRUCTURA DE ARCHIVOS
    # ----------------------------------------------------
    h1 = doc.add_heading(level=1)
    run_h1 = h1.add_run("3. Estructura de Directorios del Proyecto")
    run_h1.font.name = 'Georgia'
    run_h1.font.size = Pt(18)
    run_h1.font.bold = True
    run_h1.font.color.rgb = COLOR_PRIMARY

    p_struct = doc.add_paragraph()
    p_struct.paragraph_format.line_spacing = 1.0
    p_struct.paragraph_format.space_after = Pt(12)
    struct_code = (
        "recipes-app/\n"
        "├── index.html                   # Plantilla base del DOM y cargador HTML\n"
        "├── package.json                 # Metadatos del proyecto y dependencias de NPM\n"
        "├── vite.config.js               # Configuración de los plugins de Vite\n"
        "└── src/\n"
        "    ├── main.jsx                 # Punto de entrada de la ejecución de React\n"
        "    ├── App.jsx                  # Enrutador principal y envoltura de Context Providers\n"
        "    ├── index.css                # Estilos globales y tokens del sistema de diseño (CSS variables)\n"
        "    ├── App.css                  # Estilos auxiliares mínimos\n"
        "    ├── context/\n"
        "    │   └── FavoritesContext.jsx # Proveedor del estado de favoritos con localStorage\n"
        "    ├── hooks/\n"
        "    │   └── useDebounce.js       # Hook personalizado para optimizar las búsquedas asíncronas\n"
        "    ├── services/\n"
        "    │   └── mealApi.js           # Capa de consumo del servicio de API (TheMealDB con Axios)\n"
        "    ├── components/\n"
        "    │   ├── Navbar.jsx           # Barra de navegación superior fija y responsiva\n"
        "    │   ├── RecipeCard.jsx       # Tarjeta de receta reutilizable con hover e interactividad\n"
        "    │   ├── SkeletonCard.jsx     # Tarjeta de carga preliminar (Shimmer Effect)\n"
        "    │   ├── SearchBar.jsx        # Barra de entrada de búsqueda interactiva\n"
        "    │   ├── EmptyState.jsx       # Estado visual para listas vacías o errores de búsqueda\n"
        "    │   ├── ScrollToTop.jsx      # Auxiliar de navegación para reiniciar el scroll en cada ruta\n"
        "    │   └── Loader.jsx           # Spinner clásico para pantallas de carga fija\n"
        "    ├── pages/\n"
        "    │   ├── Home.jsx             # Vista principal con filtros, buscador y paginación avanzada\n"
        "    │   ├── Favorites.jsx        # Vista de recetas guardadas por el usuario con botón de volver\n"
        "    │   └── RecipeDetail.jsx     # Vista detallada de la receta (ingredientes, video, etc.)\n"
        "    └── styles/\n"
        "        ├── home.css             # Estilos específicos para la Home (cuadrícula, filtros, selector)\n"
        "        └── details.css          # Estilos de maquetación de la receta detallada\n"
    )
    run_sc = p_struct.add_run(struct_code)
    run_sc.font.name = 'Courier New'
    run_sc.font.size = Pt(9.5)
    run_sc.font.color.rgb = COLOR_PRIMARY

    doc.add_page_break()

    # ----------------------------------------------------
    # ANÁLISIS DETALLADO DE ARCHIVO POR ARCHIVO
    # ----------------------------------------------------
    h1 = doc.add_heading(level=1)
    run_h1 = h1.add_run("4. Análisis Detallado de Cada Archivo del Proyecto")
    run_h1.font.name = 'Georgia'
    run_h1.font.size = Pt(18)
    run_h1.font.bold = True
    run_h1.font.color.rgb = COLOR_PRIMARY
    
    doc.add_paragraph("A continuación, se describe minuciosamente la función, el flujo y el porqué de la implementación de cada archivo del proyecto, ordenados de forma lógica.")

    files_info = [
        # index.html
        {
            "filename": "index.html",
            "purpose": "Plantilla base HTML5 para el montaje del DOM de React y optimizaciones SEO.",
            "function": "1. Proporciona la estructura del documento HTML5.\n"
                        "2. Configura el idioma de la página a español (lang='es') y la codificación UTF-8.\n"
                        "3. Define el título descriptivo para SEO y añade la meta etiqueta de descripción para la previsualización en motores de búsqueda.\n"
                        "4. Contiene el div id='root' donde React monta su árbol jerárquico de componentes, y vincula el script de entrada /src/main.jsx.",
            "implements": "Maquetación HTML5 semántica y SEO básico (meta description y etiquetas de título adaptadas).",
            "works": "Sirve como el lienzo base que el navegador web descarga e inicializa antes de cargar los scripts JavaScript empaquetados por Vite.",
            "rationale": "Es el archivo de inicio estándar de cualquier SPA moderna. Al tener una única página contenedora, el DOM inicial se mantiene ligero y las optimizaciones SEO de primer nivel (como lang y description) se aplican de manera estática."
        },
        # 1. main.jsx
        {
            "filename": "src/main.jsx",
            "purpose": "Punto de entrada de la aplicación en el navegador web.",
            "function": "1. Importa la biblioteca de React y el renderizador ReactDOM.\n"
                        "2. Importa el archivo de estilos globales (index.css).\n"
                        "3. Importa el componente raíz App.jsx y la envoltura de enrutamiento BrowserRouter de react-router-dom.\n"
                        "4. Renderiza la aplicación de React de forma estricta (StrictMode) montándola sobre el nodo raíz del DOM (div id='root') del archivo index.html.",
            "implements": "Renderizado DOM en React 18, StrictMode para la detección de advertencias colaterales y BrowserRouter para habilitar enrutamiento SPA.",
            "works": "Inicializa y monta la máquina de ejecución de React en la página web.",
            "rationale": "Es la convención recomendada para aplicaciones SPA generadas con Vite, aislando la lógica de arranque de la lógica propia del negocio de los componentes."
        },
        # 2. App.jsx
        {
            "filename": "src/App.jsx",
            "purpose": "Componente raíz de la interfaz de usuario que define la estructura global y el enrutamiento.",
            "function": "1. Envuelve a toda la aplicación con el FavoritesProvider para compartir globalmente el estado de recetas favoritas.\n"
                        "2. Renderiza el componente Navbar y el componente de utilidad ScrollToTop.\n"
                        "3. Utiliza Routes y Route para mapear las URLs a las páginas correspondientes:\n"
                        "   - '/' renderiza Home.jsx\n"
                        "   - '/recipe/:id' renderiza RecipeDetail.jsx\n"
                        "   - '/favorites' renderiza Favorites.jsx\n"
                        "4. Centra todo el contenido mediante un elemento semántico <main> con clase 'main-content'.",
            "implements": "Enrutamiento dinámico SPA de react-router-dom, provisión del contexto global de favoritos y estructura base del layout.",
            "works": "Controla qué componente de nivel de página (Page) debe renderizarse en el contenedor central basándose en el pathname actual del navegador.",
            "rationale": "Mantener el enrutador en el componente raíz permite una visión global e inmediata de la estructura de URLs del sitio web, asegurando consistencia en la navegación y barras globales."
        },
        # 3. index.css
        {
            "filename": "src/index.css",
            "purpose": "Archivo global de estilos de la aplicación y declaración del sistema de diseño (tokens).",
            "function": "1. Importa las tipografías premium Montserrat (para texto e interactivos) y Playfair Display (para cabeceras de estilo revista gastronómica) desde Google Fonts.\n"
                        "2. Declara variables CSS globales en la raíz (:root) para colores (--bg, --surface, --accent, --text, --border, --shadow-sm, --shadow-md) permitiendo cambiar de tema con facilidad.\n"
                        "3. Restablece estilos por defecto, define la tipografía predeterminada y el diseño base flexible.\n"
                        "4. Contiene las reglas CSS para componentes globales: la barra de navegación superior (Navbar), los esqueletos de carga animados (SkeletonCard), el estado vacío (EmptyState), y el botón interactivo de favoritos (.recipe-card__fav-btn) con microanimaciones como @keyframes pop y .back-button.",
            "implements": "Tokens de diseño consistentes mediante Custom CSS Properties, animaciones de carga de baja latencia (@keyframes shimmer) y transiciones fluidas de hover.",
            "works": "Da forma al comportamiento visual inicial y los estilos globales de los componentes que cruzan toda la aplicación.",
            "rationale": "Permite centralizar el 'Design System' de la aplicación en un único archivo. Al definir los colores base a través de variables CSS, si en el futuro se desea cambiar la identidad visual o el color acento, solo se modifica una línea de código."
        },
        # 4. context/FavoritesContext.jsx
        {
            "filename": "src/context/FavoritesContext.jsx",
            "purpose": "Administración del estado persistente de las recetas marcadas como favoritas.",
            "function": "1. Crea un React Context llamado FavoritesContext.\n"
                        "2. Inicializa el estado 'favorites' leyendo desde localStorage mediante una función evaluada flojamente (lazy initialization).\n"
                        "3. Implementa funciones para agregar (toggleFavorite) y verificar (isFavorite) una receta favorita.\n"
                        "4. Sincroniza automáticamente cualquier cambio en la lista de favoritos con localStorage mediante un useEffect.\n"
                        "5. Exporta un hook personalizado useFavorites() que valida la correcta provisión del contexto.",
            "implements": "React Context API (createContext, useContext), persistencia nativa con LocalStorage, hooks como useState y useEffect, y validación estructurada.",
            "works": "Proporciona una interfaz global donde cualquier tarjeta de receta (RecipeCard), en cualquier parte de la aplicación, puede comprobar si es favorita o ser añadida/removida de manera inmediata, actualizando simultáneamente el contador en la Navbar.",
            "rationale": "El estado de favoritos debe ser compartido por múltiples componentes independientes (RecipeCard, Favorites page, Navbar). El uso de Context API evita el 'prop drilling' (pasar props a través de muchos niveles) y mantiene la lógica encapsulada de manera limpia."
        },
        # 5. hooks/useDebounce.js
        {
            "filename": "src/hooks/useDebounce.js",
            "purpose": "Optimizar el rendimiento y limitar la cantidad de peticiones consecutivas al buscar recetas.",
            "function": "1. Recibe un valor de texto (searchTerm) y un retardo de tiempo en milisegundos (delay).\n"
                        "2. Crea un estado interno que almacena el valor retrasado.\n"
                        "3. Configura un temporizador con setTimeout en un useEffect para actualizar el estado después del tiempo indicado.\n"
                        "4. En la fase de saneamiento (cleanup) del useEffect, limpia el temporizador en cada pulsación de tecla para evitar fugas de memoria y retrasar la ejecución si el usuario sigue escribiendo.",
            "implements": "Técnica de debouncing en React, uso de useEffect y useState, y saneamiento de temporizadores.",
            "works": "Retrasa la propagación de un cambio de estado continuo (como escribir en un input) hasta que el usuario deja de presionar teclas durante el lapso especificado.",
            "rationale": "Sin debounce, cada pulsación de tecla realizaría una solicitud HTTP inmediata a la API de recetas, saturando el servidor de red del cliente y provocando ralentizaciones (INP alto). Con debounce, solo se ejecuta la solicitud cuando el usuario termina de escribir."
        },
        # 6. services/mealApi.js
        {
            "filename": "src/services/mealApi.js",
            "purpose": "Capa de comunicación HTTP con la API pública de recetas TheMealDB.",
            "function": "1. Crea una instancia configurada de Axios con la URL base de TheMealDB.\n"
                        "2. Declara y exporta métodos específicos de consulta asíncrona:\n"
                        "   - getRecipeById(id): Solicita detalles completos de una receta por su identificador.\n"
                        "   - searchRecipes(name): Realiza búsquedas de recetas cuyo título contenga la cadena proporcionada.\n"
                        "   - getCategories(): Obtiene el listado de categorías disponibles (Dessert, Seafood, Beef, etc.).\n"
                        "   - getRecipesByCategory(category): Trae las recetas asociadas a una categoría específica.\n"
                        "   - getAllRecipes(): Trae recetas de todas las categorías en paralelo utilizando Promise.all y aplica el algoritmo de barajado Fisher-Yates para asegurar variedad en la página principal.",
            "implements": "Consumo de API REST, Axios custom instances, asincronía avanzada tolerante a fallos con async/await y Promise.all (con recuperación .catch en peticiones individuales para no cancelar la descarga completa), y mezcla de datos mediante algoritmo Fisher-Yates.",
            "works": "Se encarga de comunicarse de manera segura y controlada con la API externa, devolviendo únicamente la información útil mapeada en formato JSON.",
            "rationale": "Aislar las llamadas HTTP en un archivo de servicios es una buena práctica crítica. Si la API de recetas cambia de URL o se migra de librería en el futuro, no es necesario tocar los componentes visuales; solo se edita este archivo de servicio."
        },
        # 7. components/Navbar.jsx
        {
            "filename": "src/components/Navbar.jsx",
            "purpose": "Proporcionar la barra superior fija de navegación con el logotipo de la marca y enlaces directos.",
            "function": "1. Muestra el título de la marca ('Recetas Online') y un icono temático.\n"
                        "2. Genera enlaces reactivos a la Home ('/') y a la página de Favoritos ('/favorites').\n"
                        "3. Consume el hook useFavorites para mostrar dinámicamente un contador redondo (badge) con la cantidad exacta de recetas favoritas actualmente almacenadas.\n"
                        "4. Implementa un menú móvil desplegable (menú de hamburguesa) controlado por un estado booleano para pantallas pequeñas.",
            "implements": "Enrutamiento adaptativo NavLink de react-router-dom, menús responsive con detección de clics y consumo de estados contextuales compartidos.",
            "works": "Permite cambiar de ruta y realizar el seguimiento visual constante de los favoritos.",
            "rationale": "Debe ser interactivo y estar presente en todas las vistas de la aplicación. Mostrar el número de favoritos en tiempo real mejora la experiencia del usuario y fomenta el uso de la aplicación."
        },
        # 8. components/RecipeCard.jsx
        {
            "filename": "src/components/RecipeCard.jsx",
            "purpose": "Tarjeta modular e interactiva para representar de forma gráfica e individual una receta.",
            "function": "1. Recibe por props los datos específicos de una receta (imagen, título, categoría).\n"
                        "2. Consume el FavoritesContext a través del hook useFavorites() para saber si la receta está guardada o no.\n"
                        "3. Renderiza un botón interactivo de corazón (❤️) que flota en la parte superior derecha de la imagen.\n"
                        "4. Al hacer clic en el botón de corazón, ejecuta la acción de añadir/eliminar de favoritos de forma animada (microanimación pop).\n"
                        "5. Al pasar el cursor, ejecuta un efecto de zoom en la imagen y elevación de la tarjeta.",
            "implements": "Componentes de React controlados por props, microanimaciones CSS y reutilización visual de datos estructurados.",
            "works": "Representa de forma independiente un platillo, su categoría y permite la rápida interacción del usuario para guardarlo o abrir su ficha detallada.",
            "rationale": "La reutilización de tarjetas es fundamental en el desarrollo de software moderno. Se utiliza el mismo componente tanto en la Home como en la página de Favoritos, reduciendo la duplicación de código CSS y React."
        },
        # 9. components/SkeletonCard.jsx
        {
            "filename": "src/components/SkeletonCard.jsx",
            "purpose": "Evitar cambios bruscos en el diseño (Content Layout Shift) mostrando una estructura de carga simulada.",
            "function": "1. Renderiza bloques de div planos con la misma relación de aspecto y dimensiones que una tarjeta RecipeCard real.\n"
                        "2. Aplica la clase '.skeleton' la cual contiene un efecto de degradado metálico en movimiento constante animado por keyframes (shimmer effect).\n"
                        "3. Sirve para avisar al usuario de que el contenido está en camino sin bloquear la interfaz.",
            "implements": "Patrones modernos de carga UX, animaciones basadas en CSS de baja carga computacional y prevención de CLS.",
            "works": "Se dibuja en la pantalla de forma temporal mientras la llamada HTTP asíncrona se resuelve de fondo en la Home.",
            "rationale": "Mostrar un loader tradicional o un simple mensaje de 'Cargando...' resulta en una experiencia de usuario desactualizada. Los esqueletos mantienen la estructura visual intacta, reduciendo la impaciencia percibida del usuario."
        },
        # 10. components/EmptyState.jsx
        {
            "filename": "src/components/EmptyState.jsx",
            "purpose": "Mostrar un mensaje visual amigable e interactivo cuando una pantalla o cuadrícula de búsqueda no tiene datos.",
            "function": "1. Recibe mediante props un icono (emoji), un título llamativo y un mensaje explicativo.\n"
                        "2. Renderiza un diseño centrado con tipografía de calidad y colores suavizados que guían al usuario sobre qué hacer a continuación.",
            "implements": "Representación de estados alternativos (Empty States) y tipografía editorial en componentes planos.",
            "works": "Se activa dinámicamente cuando un término de búsqueda no coincide con ninguna receta o cuando no hay elementos en favoritos.",
            "rationale": "Evita dejar al usuario con una pantalla en blanco y proporciona instrucciones claras para continuar navegando."
        },
        # 11. components/SearchBar.jsx
        {
            "filename": "src/components/SearchBar.jsx",
            "purpose": "Input de búsqueda interactivo que actúa como el primer canal de entrada para filtrar recetas por texto.",
            "function": "1. Recibe el estado searchTerm y su setter setSearchTerm mediante props.\n"
                        "2. Renderiza un campo de entrada <input> limpio, con un placeholder descriptivo y un diseño de bordes redondeados y sombreado suave.",
            "implements": "Entradas controladas en React (controlled components) y estilos de foco dinámicos.",
            "works": "Actualiza el estado de búsqueda en el componente padre inmediato (Home) cada vez que el usuario teclea un caracter.",
            "rationale": "Encapsular el input permite reutilizar o reestilizar la barra de búsqueda en otras secciones si fuese necesario, manteniendo limpio el código de la página Home."
        },
        # 12. components/ScrollToTop.jsx
        {
            "filename": "src/components/ScrollToTop.jsx",
            "purpose": "Corregir el comportamiento nativo del scroll del navegador al realizar enrutamientos dinámicos.",
            "function": "1. Utiliza el hook useLocation para enterarse del cambio de la ruta actual del sitio web.\n"
                        "2. Ejecuta un useEffect que llama a window.scrollTo(0, 0) cada vez que el pathname de navegación cambia.\n"
                        "3. No renderiza ningún elemento visual (retorna null).",
            "implements": "Hooks de ciclo de vida en enrutamiento (useLocation) e interacción con las APIs del navegador (Window API).",
            "works": "Fuerza de manera silenciosa al navegador a colocarse en la coordenada vertical cero cada vez que cargamos una página.",
            "rationale": "En aplicaciones SPA, si un usuario hace scroll hacia abajo en una página larga y luego hace clic en un enlace a otra página, el navegador mantiene la posición de scroll baja en la nueva ruta. Este componente corrige esta deficiencia de usabilidad."
        },
        # 13. components/Loader.jsx
        {
            "filename": "src/components/Loader.jsx",
            "purpose": "Componente spinner clásico utilizado como cargador alternativo en vistas específicas.",
            "function": "1. Renderiza un contenedor centrado con un elemento de carga animado de forma giratoria continua por CSS.\n"
                        "2. Se muestra durante el tiempo de resolución de recetas individuales.",
            "implements": "Animaciones giratorias CSS a través de @keyframes spin y maquetación centrada.",
            "works": "Representa el estado activo de una petición asíncrona de manera visual clásica.",
            "rationale": "Proporciona una solución ligera para componentes secundarios o pantallas donde no se previsualiza una cuadrícula de tarjetas completas."
        },
        # 14. pages/Home.jsx
        {
            "filename": "src/pages/Home.jsx",
            "purpose": "Pantalla de inicio principal de la aplicación.",
            "function": "1. Declara estados clave: lista de recetas (allRecipes), término de búsqueda (searchTerm), seleccionado de categorías (selectedCategory), página actual (currentPage), carga (loading) y error (error).\n"
                        "2. Incorpora el hook useDebounce para filtrar las peticiones asíncronas de búsqueda de texto.\n"
                        "3. Lanza solicitudes HTTP asíncronas a mealApi basadas en el filtro o búsqueda activa a través de un useEffect coordinado.\n"
                        "4. Implementa una rejilla adaptativa de tarjetas y renderiza SkeletonCard mientras carga.\n"
                        "5. Gestiona una paginación inteligente de 12 recetas por página mediante un algoritmo de ventana deslizante que evita que los botones desaparezcan bruscamente.",
            "implements": "Gestión de estados interconectados en React, ciclos de vida dinámicos (useEffect) optimizados con bandera de cancelación (active flag) para evitar condiciones de carrera, renderizados condicionales múltiples y algoritmos de paginación fluida.",
            "works": "Centraliza la experiencia principal del usuario permitiendo buscar, filtrar por categorías y navegar por páginas de resultados culinarios con total dinamismo.",
            "rationale": "Es el cerebro de la aplicación. Al gestionar los filtros y la búsqueda de forma centralizada en este archivo, se asegura que las recetas se actualicen de manera síncrona y fluida en respuesta a cualquier interacción del usuario."
        },
        # 15. pages/Favorites.jsx
        {
            "filename": "src/pages/Favorites.jsx",
            "purpose": "Pantalla dedicada a renderizar las recetas guardadas por el usuario como favoritas.",
            "function": "1. Consume el contexto de favoritos mediante el hook useFavorites().\n"
                        "2. Si la lista está vacía, renderiza el componente EmptyState con el icono interactivo del corazón roto.\n"
                        "3. Si tiene recetas favoritas, renderiza una cuadrícula (recipes-grid) con las recetas correspondientes.\n"
                        "4. Añade un botón interactivo de regreso a la Home ('Volver a recetas') en la parte superior que respeta la coherencia del flujo de navegación.",
            "implements": "Consumo de estados del FavoritesContext, renderizado dinámico condicional y enrutamiento inverso.",
            "works": "Extrae las recetas almacenadas localmente del navegador del usuario y las previsualiza de la misma manera refinada que la página de inicio.",
            "rationale": "Ofrece un espacio privado de recopilación para el usuario. Reutiliza los estilos globales y las tarjetas para mantener una estética consistente en todo el portal culinario."
        },
        # 16. pages/RecipeDetail.jsx
        {
            "filename": "src/pages/RecipeDetail.jsx",
            "purpose": "Pantalla detallada para mostrar el contenido completo de una sola receta culinaria.",
            "function": "1. Captura el parámetro de identificación (id) de la URL usando useParams().\n"
                        "2. Realiza la llamada HTTP asíncrona getRecipeById(id) en un useEffect protegido contra condiciones de carrera mediante bandera de cancelación (active flag).\n"
                        "3. Extrae dinámicamente hasta 20 ingredientes y sus respectivas medidas del objeto meal devuelto por la API mediante un bucle iterativo.\n"
                        "4. Analiza de manera robusta la URL de YouTube utilizando una función utilitaria con expresión regular que contempla múltiples formatos de enlace (cortos, estándar, embebidos) para extraer el ID correcto del video.\n"
                        "5. Muestra el botón para regresar a la página principal de recetas.",
            "implements": "Enrutamiento paramétrico (useParams), manejo de efectos con cancelación, expresiones regulares avanzadas de análisis, renderizado dinámico de listas asociativas, e iframe interactivo responsivo para reproducción de video.",
            "works": "Consulta la información detallada del platillo al servidor y compone una vista de revista gastronómica dividida en secciones lógicas (Ingredientes, Instrucciones, Video).",
            "rationale": "Ofrece el nivel más profundo de interacción del usuario. Extraer los ingredientes vacíos dinámicamente en el cliente limpia la información devuelta por la base de datos de TheMealDB, asegurando que solo se renderice lo que se va a usar."
        },
        # 17. styles/home.css
        {
            "filename": "src/styles/home.css",
            "purpose": "Estilos específicos para la Home Page y la maquetación de la cuadrícula.",
            "function": "1. Estiliza el título h1 con una línea naranja inferior que refuerza la estética editorial.\n"
                        "2. Estiliza el contenedor de filtros, el input de búsqueda y el menú select de categorías con un aspecto limpio y bordes sutiles.\n"
                        "3. Diseña la cuadrícula de recetas (recipes-grid) con CSS Grid adaptativo y las tarjetas individuales con sombras profundas y transiciones suaves.\n"
                        "4. Controla la visualización del Loader y los mensajes de error.\n"
                        "5. Contiene los estilos para el componente de navegación de paginación (.pagination y .pagination-btn) con un look moderno y adaptativo en móviles.",
            "implements": "Diseño web adaptable (Responsive Web Design), CSS Grid Layout y transiciones de aceleración por hardware.",
            "works": "Define la identidad visual concreta y el ordenamiento de los elementos interactivos en la pantalla de inicio.",
            "rationale": "Al separar la apariencia de la Home de la página de detalles, se garantiza que los archivos de estilos sean ligeros, limpios y fáciles de depurar en caso de variaciones del layout."
        },
        # 18. styles/details.css
        {
            "filename": "src/styles/details.css",
            "purpose": "Estilos de maquetación detallada para la ficha individual del platillo.",
            "function": "1. Estiliza el contenedor de detalles con anchos máximos controlados y márgenes elegantes.\n"
                        "2. Aplica estilos premium a las listas de ingredientes (ingredients-list) distribuyéndolas en rejillas de dos columnas en pantallas de escritorio.\n"
                        "3. Estiliza el texto largo de las instrucciones con una capitular elegante en la primera letra (.detail-container > p::first-letter) emulando revistas gourmet.\n"
                        "4. Garantiza que la reproducción de videos embebidos de YouTube sea responsiva y mantenga la proporción 16:9 de forma fluida.",
            "implements": "Capitulares en tipografía CSS, cuadrículas flexibles y adaptabilidad de medios embebidos.",
            "works": "Distribuye visualmente la información detallada para que sea fácil de leer en cualquier dispositivo móvil o de escritorio.",
            "rationale": "Facilita la lectura de textos largos e instrucciones complejas mientras el usuario cocina, usando una jerarquía tipográfica amplia y espaciados generosos."
        }
    ]

    # Add each file to the document with structured subheadings
    for file_info in files_info:
        h2 = doc.add_heading(level=2)
        run_h2 = h2.add_run(f"📂 Archivo: {file_info['filename']}")
        run_h2.font.name = 'Georgia'
        run_h2.font.size = Pt(14)
        run_h2.font.bold = True
        run_h2.font.color.rgb = COLOR_SECONDARY
        
        # Create table for detailed analysis (Clean Word styling)
        table = doc.add_table(rows=5, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(1.8)
        table.columns[1].width = Inches(4.7)
        
        headers = [
            ("¿Para qué sirve?", file_info['purpose']),
            ("Función paso a paso", file_info['function']),
            ("¿Qué se implementa?", file_info['implements']),
            ("¿Cómo funciona?", file_info['works']),
            ("¿Por qué se hizo así?", file_info['rationale'])
        ]
        
        for idx, (label, content) in enumerate(headers):
            row = table.rows[idx]
            cell_label = row.cells[0]
            cell_content = row.cells[1]
            
            # Label Styling
            p_label = cell_label.paragraphs[0]
            p_label.paragraph_format.space_after = Pt(4)
            run_lbl = p_label.add_run(label)
            run_lbl.font.bold = True
            run_lbl.font.size = Pt(10)
            run_lbl.font.color.rgb = COLOR_PRIMARY
            set_cell_background(cell_label, "F4F5F7")
            set_cell_margins(cell_label, top=80, bottom=80, left=100, right=100)
            
            # Content Styling
            p_content = cell_content.paragraphs[0]
            p_content.paragraph_format.line_spacing = 1.15
            p_content.paragraph_format.space_after = Pt(4)
            run_cnt = p_content.add_run(content)
            run_cnt.font.size = Pt(10)
            run_cnt.font.color.rgb = COLOR_TEXT
            set_cell_margins(cell_content, top=80, bottom=80, left=100, right=100)
            
        doc.add_paragraph() # Spacer

    doc.add_page_break()

    # ----------------------------------------------------
    # DISEÑO Y ESTÉTICA PREMIUM (Gourmet UX)
    # ----------------------------------------------------
    h1 = doc.add_heading(level=1)
    run_h1 = h1.add_run("5. Principios de UX y Estética Visual Implementados")
    run_h1.font.name = 'Georgia'
    run_h1.font.size = Pt(18)
    run_h1.font.bold = True
    run_h1.font.color.rgb = COLOR_PRIMARY
    
    p = doc.add_paragraph(
        "El diseño de la aplicación no es genérico. Se concibió como un portal gastronómico "
        "editorial de primer nivel, cuidando detalles tipográficos y de interacción para "
        "ofrecer una experiencia premium:"
    )
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(12)

    ux_points = [
        ("Tipografía Asimétrica y Jerárquica: ", "El uso de la fuente Playfair Display (Serif) con anchos gruesos y bordes definidos le otorga un estilo editorial tipo revista de cocina a los títulos principales, mientras que Montserrat (Sans-Serif) asegura legibilidad en los ingredientes e instrucciones largas."),
        ("Paleta de Colores de Alta Gama: ", "Se evitó el uso de colores planos o chillones. Se adoptó una paleta basada en gris oscuro carbón (#1A1B1E) para textos y un color ámbar/terracota (#D97706) para acentos que evoca ingredientes frescos, calor de hogar y sofisticación."),
        ("Elevación Visual y Hover Micro-interactivo: ", "Las tarjetas de recetas no son estáticas. Poseen transiciones suaves de zoom en imagen y elevación (`translateY(-6px)`) con sombras difusas, invitando físicamente al usuario a interactuar con ellas."),
        ("Prevención del Parpadeo de Carga (Skeletons): ", "La inclusión de Skeleton Cards con efecto 'shimmer' animado simula la presencia física del contenido. Esto disminuye la tasa de rebote del usuario y suaviza la transición de los datos."),
        ("Enrutamiento Fluido y Scroll Automático: ", "Gracias a ScrollToTop y React Router, la navegación simula transiciones instantáneas donde el usuario no tiene que reiniciar de forma manual la vista de scroll vertical cada vez que navega a un platillo."),
        ("Paginación Robusta de Ventana Deslizable: ", "El algoritmo de paginación reajusta dinámicamente el rango de páginas visibles, garantizando que los usuarios puedan moverse de forma lineal o saltar de extremo a extremo sin que la botonera se desorganice visualmente en pantallas grandes o móviles.")
    ]

    for title, desc in ux_points:
        p_u = doc.add_paragraph(style='List Bullet')
        p_u.paragraph_format.space_after = Pt(6)
        r_title = p_u.add_run(title)
        r_title.font.bold = True
        r_title.font.color.rgb = COLOR_SECONDARY
        r_desc = p_u.add_run(desc)
        r_desc.font.color.rgb = COLOR_TEXT

    # Save document
    output_filename = "DOCUMENTACION_TECNICA.docx"
    doc.save(output_filename)
    print(f"Documento '{output_filename}' generado con éxito.")

if __name__ == "__main__":
    generate_document()
